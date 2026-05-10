"""
Vet Certificate Generator v2 — El-Anam Veterinary Clinic
=========================================================
• Clinic selector: خميس مشيط / ظبية — switches stamp, trade reg, license
• Stamp & signature read from: D:\\Anaam\\certifcate_generator\\
• Chip or Ring number toggle (per-animal)
• Species manager panel (add/remove scientific names, persisted to JSON)
• Word (.docx) export checkbox
• Output folder: D:\\Anaam\\certifcate_generator\\outputs\\
• File named: <owner>_<species>_<type>.pdf|.docx
• Professional ttkbootstrap UI
"""

from __future__ import annotations
import os, json, random
from datetime import datetime
from dataclasses import dataclass
from typing import List

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import ttkbootstrap as tb
from ttkbootstrap.constants import *

from reportlab.lib.pagesizes import A4
from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                Paragraph, Spacer, Image as RLImage)
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import mm

from PIL import Image

# QR code generation
try:
    import qrcode
    _qrcode_available = True
except ImportError:
    _qrcode_available = False
    print("[WARN] qrcode not installed. Install with: pip install qrcode[pil]")

# Arabic text rendering: reshape + bidi reorder (BOTH are required for correct Arabic in PDF)
try:
    import arabic_reshaper
    _reshaper_available = True
except ImportError:
    _reshaper_available = False
    print("[WARN] arabic_reshaper not installed. Install with: pip install arabic-reshaper")

try:
    from bidi.algorithm import get_display
    _bidi_available = True
except ImportError:
    _bidi_available = False
    print("[WARN] bidi library not installed. Install with: pip install python-bidi")

# ─────────────────────────────────────────────
#  PATHS
# ─────────────────────────────────────────────
BASE_DIR    = r"D:\mahmoud_analysis_dashboard_folder\HEALTH_CERTIFCATE\New folder"
OUT_DIR      = os.path.join(BASE_DIR, "outputs")
OUT_ISDAR    = os.path.join(BASE_DIR, "output_isdar")
SPECIES_JSON = os.path.join(BASE_DIR, "species_list.json")



# BASE_DIR    = r"D:\mahmoud_analysis_dashboard_folder\HEALTH_CERTIFCATE\New folder"
# OUT_DIR     = os.path.join(BASE_DIR, "outputs")
# SPECIES_JSON = os.path.join(BASE_DIR, "species_list.json")
#
# # Font priority: Amiri (best) > Traditional Arabic Bold > Tahoma > Arial Unicode
_FONT_CANDIDATES = [
    r"C:\Windows\Fonts\amiri-regular.ttf",
    r"C:\Windows\Fonts\Amiri-Regular.ttf",
    r"C:\Windows\Fonts\tradbdo.ttf",
    r"C:\Windows\Fonts\tahoma.ttf",
    r"C:\Windows\Fonts\arialuni.ttf",
    r"C:\Windows\Fonts\arial.ttf",
]
FONT_PATH = None
for _fp in _FONT_CANDIDATES:
    if os.path.exists(_fp):
        FONT_PATH = _fp
        print(f"[FONT] Using: {_fp}")
        break
if not FONT_PATH:
    print("[WARN] No Arabic font found. Download Amiri from https://fonts.google.com/specimen/Amiri")
#
# # Stamp & signature images — inside BASE_DIR
# STAMP_KHAMIS = os.path.join(BASE_DIR, "stamp_khamis.png")
# STAMP_ZAPIA  = os.path.join(BASE_DIR, "stamp_zapia.png")
# SIG_1        = os.path.join(BASE_DIR, "signature_1.png")
# SIG_2        = os.path.join(BASE_DIR, "signature_2.png")
# SIG_3        = os.path.join(BASE_DIR, "signature_3.png")
#
# # Header / Logo (shared)
# LOGO_PNG     = os.path.join(BASE_DIR, "logo.png")   # clinic logo top-right
#
# os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(OUT_ISDAR, exist_ok=True)

























# Try multiple possible font paths (Amiri or Arial)
# FONT_PATHS = [
#        r"C:/Windows/Fonts/arial.ttf",
#     r"C:/Windows/Fonts/arialuni.ttf",
#     r"C:/Windows/Fonts/tahoma.ttf"
# ]
# FONT_PATH = None
# for fp in FONT_PATHS:
#     if os.path.exists(fp):
#         FONT_PATH = fp
#         break
# if not FONT_PATH:
#     # fallback to default Helvetica (will not show Arabic correctly)
#     FONT_PATH = None



FONT_PATH = r"C:\Windows\Fonts\tradbdo.ttf"

# Stamp & signature images — inside BASE_DIR
STAMP_KHAMIS = os.path.join(BASE_DIR, "stamp_khamis.png")
STAMP_ZAPIA  = os.path.join(BASE_DIR, "stamp_zapia.png")
SIG_1        = os.path.join(BASE_DIR, "signature_1.png")
SIG_2        = os.path.join(BASE_DIR, "signature_2.png")
SIG_3        = os.path.join(BASE_DIR, "signature_3.png")

# Header / Logo (shared)
LOGO_PNG     = os.path.join(BASE_DIR, "logo.png")

os.makedirs(OUT_DIR, exist_ok=True)














# ─────────────────────────────────────────────
#  CLINIC DATA
# ─────────────────────────────────────────────
CLINICS = {
    "خميس مشيط": {
        "name_ar":       "عيادة الأنعام الأليفة – خميس مشيط",
        "trade_reg":     "5855357377",
        "tax_reg":       "310499782200003",
        "license":       "22/06/750012/1125",
        "address_ar":    "خميس مشيط – حي العرق الشمالي -قمبر",
        "stamp_path":    STAMP_KHAMIS,
        "phone":         "0559164099",
    },
    "ظبية": {
        "name_ar":       "عيادة الأنعام الأليفة – ظبية",
        "trade_reg":     "5900139881",
        "tax_reg":       "310499782200003",
        "license":       "23/10/750012/065",
        "address_ar":    "ظبية",
        "stamp_path":    STAMP_ZAPIA,
        "phone":         "0559164099",
    },
}

CLINIC_LOGO_AR = "عيادة الأنعام الأليفة البيطرية"
# Paragraph template with phone placeholder
PARAGRAPH_TEMPLATE = (
    "ترى عيادة الأنعام الأليفة البيطرية بأن الحيوانات المذكورة مواصفتها "
    "أدناه سليمة صحاً ظاهراً وليس بها أي أعراض مرضية. "
    "رقم التواصل مع العيادة {}"
)

# ─────────────────────────────────────────────
#  DEFAULT SPECIES LIST
# ─────────────────────────────────────────────
DEFAULT_SPECIES = [
    "Capra nubiana", "Gazella Arabica", "Gazella gazella",
    "Alectoris melanocephala", "Alectoris chukar", "Alectoris philbyi",
    "Ammoperdix heyi", "Procavia Capensis", "Lepus capensis",
    "Ammopredix griseogularis", "Gazella Marica", "Dolichatis patagonum",
    "Oryx leucoryx", "Struthio camelus", "Dromaius novaehollandiae",
    "Streptopelia turtur", "Pterocles alchata", "Chlamydotis macqueenii",
    "Anthropoides virgo", "Oryx dammah", "Addax nasomaculatus",
    "Vicugna pacos", "Ovis gmelini", "Equus zebra zebra",
    "Hystrix indica", "Phasianus colchicus", "Chrysolophus pictus",
    "Procyon totor",
]

def load_species() -> List[str]:
    if os.path.exists(SPECIES_JSON):
        try:
            with open(SPECIES_JSON, encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list) and data:
                return sorted(set(data))
        except Exception:
            pass
    return sorted(set(DEFAULT_SPECIES))

def save_species(lst: List[str]):
    try:
        with open(SPECIES_JSON, "w", encoding="utf-8") as f:
            json.dump(sorted(set(lst)), f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"[WARN] Could not save species list: {e}")

# ─────────────────────────────────────────────
#  ARABIC FONT & RENDERING
# ─────────────────────────────────────────────
AR_FONT = "Helvetica"  # fallback
if FONT_PATH:
    try:
        pdfmetrics.registerFont(TTFont("ArabicFont", FONT_PATH))
        AR_FONT = "ArabicFont"
    except Exception as e:
        print(f"[WARN] Could not register font {FONT_PATH}: {e}")

def ar(txt: str) -> str:
    """
    Correct Arabic rendering for ReportLab PDF:
    Step 1 - arabic_reshaper: connects letters and applies correct letter forms (ligatures).
    Step 2 - bidi get_display: reorders characters right-to-left.
    Both steps are REQUIRED. Skipping reshaper causes disconnected/wrong letter shapes.
    """
    if not txt:
        return ""
    try:
        # Step 1: reshape (connect Arabic letters properly)
        if _reshaper_available:
            reshaped = arabic_reshaper.reshape(txt)
        else:
            reshaped = txt  # will look broken without reshaper

        # Step 2: bidi reorder (right-to-left direction)
        if _bidi_available:
            return get_display(reshaped)
        else:
            return reshaped[::-1]  # crude fallback
    except Exception as e:
        print(f"[ar()] Error processing text: {e}")
        return txt

# ─────────────────────────────────────────────
#  DATA CLASSES
# ─────────────────────────────────────────────
@dataclass
class AnimalRow:
    species:    str
    age:        str
    id_number:  str
    id_type:    str   # "شريحة" or "حلقة"
    gender:     str
    count:      int = 1

# ─────────────────────────────────────────────
#  PDF BUILDER
# ─────────────────────────────────────────────
def _safe_img(path: str, max_w: float, max_h: float) -> Optional[RLImage]:
    if not path or not os.path.exists(path):
        return None
    try:
        with Image.open(path) as img:
            iw, ih = img.size
        iw_pt = iw * 72 / 96
        ih_pt = ih * 72 / 96
        scale = min(max_w / iw_pt, max_h / ih_pt, 1.5)
        return RLImage(path, width=iw_pt * scale, height=ih_pt * scale)
    except Exception as e:
        print(f"[IMG] {path}: {e}")
        return None

def _make_qr(owner: str, owner_id: str, species_list: List[AnimalRow], date: str = "", serial: str = "") -> Optional[RLImage]:
    """Generate a small QR code with owner + animal info and return as RLImage."""
    if not _qrcode_available:
        return None
    try:
        import io
        species_names = ", ".join(a.species for a in species_list)
        qr_data = f"المالك: {owner} | الهوية: {owner_id} | الحيوان: {species_names} | التاريخ: {date} | رقم الإصدار: {serial}"
        qr = qrcode.QRCode(version=2, error_correction=qrcode.constants.ERROR_CORRECT_M,
                           box_size=4, border=2)
        qr.add_data(qr_data)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        size = 22 * mm  # small — about 2.2 cm
        return RLImage(buf, width=size, height=size)
    except Exception as e:
        print(f"[QR] Could not generate QR: {e}")
        return None


def build_pdf(
    clinic_key: str,
    owner: str,
    owner_id: str,
    date: str,
    serial: str,
    animals: List[AnimalRow],
    out_path: str,
):
    clinic = CLINICS[clinic_key]
    page_w, page_h = A4
    lm = rm = 14 * mm
    avail_w = page_w - lm - rm

    tm = bm = 7 * mm   # tighter top/bottom margins
    doc = SimpleDocTemplate(
        out_path, pagesize=A4,
        leftMargin=lm, rightMargin=rm, topMargin=tm, bottomMargin=bm,
    )

    # styles
    def ps(name, **kw):
        defaults = dict(fontName=AR_FONT, fontSize=11, leading=15, spaceAfter=2)
        defaults.update(kw)
        return ParagraphStyle(name, **defaults)

    s_center   = ps("ctr",  alignment=TA_CENTER)
    s_right    = ps("rt",   alignment=TA_RIGHT, fontName=AR_FONT, fontSize=11, leading=16)
    s_red_rt   = ps("redrt",alignment=TA_RIGHT, textColor=colors.HexColor("#c0392b"), fontSize=12, fontName=AR_FONT, leading=17)
    s_body     = ps("body", alignment=TA_RIGHT, fontSize=11, leading=17, fontName=AR_FONT)
    s_label    = ps("lbl",  fontSize=10, leading=13, alignment=TA_RIGHT, fontName=AR_FONT)
    s_value    = ps("val",  fontSize=10, leading=13, alignment=TA_RIGHT, fontName=AR_FONT)

    els = []

    # Header
    # ── Header styles ──
    s_clinic_name = ps("hname",
        fontName=AR_FONT, fontSize=15, leading=22,
        alignment=TA_RIGHT, textColor=colors.HexColor("#1a5276"))
    s_hinfo = ps("hinfo",
        fontName=AR_FONT, fontSize=10, leading=16,
        alignment=TA_RIGHT, textColor=colors.HexColor("#2c3e50"),
        spaceAfter=0)

    logo = _safe_img(LOGO_PNG, 50*mm, 38*mm)
    logo_cell = logo if logo else Paragraph(ar(CLINIC_LOGO_AR), s_center)

    # Header right column: name + two info lines with equal spacing
    header_info = Table(
        [
            [Paragraph(ar(clinic["name_ar"]), s_clinic_name)],
            [Paragraph(clinic["trade_reg"] + "  " + ar("سجل تجاري :"), s_hinfo)],
            [Paragraph(clinic["tax_reg"]   + "  " + ar("رقم السجل الضريبي :"), s_hinfo)],
        ],
        colWidths=[avail_w - 54*mm],
        rowHeights=[None, None, None],
    )
    header_info.setStyle(TableStyle([
        ("ALIGN",         (0,0), (-1,-1), "RIGHT"),
        ("VALIGN",        (0,0), (-1,-1), "TOP"),
        ("TOPPADDING",    (0,0), (-1,-1), 2),
        ("BOTTOMPADDING", (0,0), (-1,-1), 2),
        ("LEFTPADDING",   (0,0), (-1,-1), 0),
        ("RIGHTPADDING",  (0,0), (-1,-1), 0),
    ]))

    hdr_tbl = Table(
        [[logo_cell, header_info]],
        colWidths=[54*mm, avail_w - 54*mm],
    )
    hdr_tbl.setStyle(TableStyle([
        ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
        ("ALIGN",         (0,0), (0,0),   "LEFT"),
        ("ALIGN",         (1,0), (1,0),   "RIGHT"),
        ("LEFTPADDING",   (0,0), (-1,-1), 0),
        ("RIGHTPADDING",  (0,0), (-1,-1), 0),
        ("TOPPADDING",    (0,0), (-1,-1), 0),
        ("BOTTOMPADDING", (0,0), (-1,-1), 0),
    ]))
    els.append(hdr_tbl)
    els.append(Spacer(1, 2*mm))

    # Divider
    div = Table([[""]], colWidths=[avail_w])
    div.setStyle(TableStyle([
        ("LINEBELOW", (0,0), (-1,-1), 1.5, colors.HexColor("#2980b9")),
        ("BOTTOMPADDING", (0,0), (-1,-1), 0),
        ("TOPPADDING",    (0,0), (-1,-1), 0),
    ]))
    els.append(div)
    els.append(Spacer(1, 2*mm))

    # Owner info — each field on its own line, label + value right-aligned
    owner_disp = ar(owner) if any('\u0600'<=c<='\u06FF' for c in owner) else owner

    s_lbl2  = ps("lbl2",  fontSize=11, leading=17, alignment=TA_RIGHT,
                 fontName=AR_FONT, textColor=colors.HexColor("#2c3e50"))
    s_val2  = ps("val2",  fontSize=11, leading=17, alignment=TA_RIGHT,
                 fontName=AR_FONT, textColor=colors.HexColor("#2c3e50"))
    s_date2 = ps("dat2",  fontSize=12, leading=17, alignment=TA_RIGHT,
                 fontName=AR_FONT, textColor=colors.HexColor("#c0392b"))

    # Label width | Value width  (two columns, right-to-left reading)
    lw = 32 * mm   # label column
    vw = avail_w - lw

    owner_rows = [
        [Paragraph(owner_disp, s_val2),         Paragraph(ar("اسم المالك :"), s_lbl2)],
        [Paragraph(owner_id,    s_val2),         Paragraph(ar("رقم الهوية :"), s_lbl2)],
        [Paragraph(f"<b>{date}</b>", s_date2),   Paragraph(ar("تاريخ الإصدار :"), s_lbl2)],
    ]
    oi_tbl = Table(owner_rows, colWidths=[vw, lw])
    oi_tbl.setStyle(TableStyle([
        ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
        ("ALIGN",         (0,0), (-1,-1), "RIGHT"),
        ("TOPPADDING",    (0,0), (-1,-1), 2),
        ("BOTTOMPADDING", (0,0), (-1,-1), 2),
        ("LEFTPADDING",   (0,0), (-1,-1), 2),
        ("RIGHTPADDING",  (0,0), (-1,-1), 2),
    ]))
    els.append(oi_tbl)
    els.append(Spacer(1, 3*mm))

    # Body paragraph: reshape+bidi the Arabic part, keep phone number as plain suffix
    paragraph_arabic = ar(
        "ترى عيادة الأنعام الأليفة البيطرية بأن الحيوانات المذكورة مواصفتها "
        "أدناه سليمة صحاً ظاهراً وليس بها أي أعراض مرضية. "
        "رقم التواصل مع العيادة"
    )
    paragraph_text = paragraph_arabic + " " + clinic["phone"]
    els.append(Paragraph(paragraph_text, s_body))
    els.append(Spacer(1, 3*mm))

    # Animals table — dynamic header based on actual id_type used
    id_types_used = list(dict.fromkeys(a.id_type for a in animals))  # unique, ordered
    if len(id_types_used) == 1:
        id_header_text = f"رقم ال{id_types_used[0]}"
    else:
        id_header_text = "رقم الشريحة / الحلقة"
    id_header = ar(id_header_text)
    headers = [ar("العدد"), ar("الجنس"), id_header, ar("العمر"), ar("الاسم العلمي")]
    tbl_data = [headers]
    for a in animals:
        age_txt = ar(a.age)  # age already includes unit
        tbl_data.append([
            str(a.count),
            ar(a.gender),
            str(a.id_number),
            age_txt,
            a.species,
        ])

    col_w2 = [12*mm, 18*mm, 48*mm, 22*mm, avail_w - 12*mm - 18*mm - 48*mm - 22*mm]
    at = Table(tbl_data, colWidths=col_w2, repeatRows=1,
               rowHeights=[9*mm] + [8.5*mm]*(len(tbl_data)-1))
    at.setStyle(TableStyle([
        ("BACKGROUND",   (0,0), (-1,0),  colors.HexColor("#2980b9")),
        ("TEXTCOLOR",    (0,0), (-1,0),  colors.white),
        ("FONTNAME",     (0,0), (-1,0),  AR_FONT),
        ("FONTSIZE",     (0,0), (-1,0),  10),
        ("FONTNAME",     (0,1), (-1,-1), AR_FONT),
        ("FONTSIZE",     (0,1), (-1,-1), 9),
        ("TEXTCOLOR",    (0,1), (-1,-1), colors.HexColor("#2c3e50")),
        ("ALIGN",        (0,0), (-1,-1), "CENTER"),
        ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
        ("GRID",         (0,0), (-1,-1), 0.5, colors.HexColor("#bdc3c7")),
        ("BOX",          (0,0), (-1,-1), 1.2, colors.HexColor("#2c3e50")),
        ("LEFTPADDING",  (0,0), (-1,-1), 3),
        ("RIGHTPADDING", (0,0), (-1,-1), 3),
        ("TOPPADDING",   (0,0), (-1,-1), 2),
        ("BOTTOMPADDING",(0,0), (-1,-1), 2),
    ]))
    for i in range(2, len(tbl_data), 2):
        at.setStyle(TableStyle([("BACKGROUND", (0,i), (-1,i), colors.HexColor("#eaf4fb"))]))
    els.append(at)
    els.append(Spacer(1, 3*mm))

    # Serial, address, license — RIGHT side | QR code — LEFT side
    qr_img = _make_qr(owner, owner_id, animals, date, serial)
    qr_cell = qr_img if qr_img else Paragraph("", s_center)

    info_lines = [
        Paragraph("<b>" + serial + "</b>  " + ar("رقم الإصدار :"), s_red_rt),
        Paragraph(ar(clinic["address_ar"]), s_right),
        Paragraph(clinic["license"] + "  " + ar("ترخيص العيادة :"), s_right),
    ]

    info_tbl = Table(
        [[qr_cell, info_lines]],
        colWidths=[26*mm, avail_w - 26*mm],
    )
    info_tbl.setStyle(TableStyle([
        ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
        ("ALIGN",         (0,0), (0,0),   "LEFT"),
        ("ALIGN",         (1,0), (1,0),   "RIGHT"),
        ("LEFTPADDING",   (0,0), (-1,-1), 0),
        ("RIGHTPADDING",  (0,0), (-1,-1), 0),
    ]))
    els.append(info_tbl)
    els.append(Spacer(1, 4*mm))

    # Stamp and signature
    stamp_img = _safe_img(clinic["stamp_path"], 55*mm, 48*mm)
    sig_img   = _safe_img(SIG_1, 52*mm, 40*mm)
    stamp_cell = stamp_img if stamp_img else Paragraph(ar("الختم"), s_center)
    sig_cell   = sig_img   if sig_img   else Paragraph(ar("التوقيع"), s_center)

    footer_tbl = Table(
        [[stamp_cell, "", sig_cell]],
        colWidths=[58*mm, avail_w - 116*mm, 58*mm],
    )
    footer_tbl.setStyle(TableStyle([
        ("ALIGN",  (0,0), (0,0), "LEFT"),
        ("ALIGN",  (2,0), (2,0), "RIGHT"),
        ("VALIGN", (0,0), (-1,-1), "BOTTOM"),
        ("LEFTPADDING",  (0,0), (-1,-1), 0),
        ("RIGHTPADDING", (0,0), (-1,-1), 0),
    ]))
    els.append(footer_tbl)

    doc.build(els)

# ─────────────────────────────────────────────
#  WORD EXPORT (with images and larger signature)
# ─────────────────────────────────────────────
def build_docx(
    clinic_key: str,
    owner: str,
    owner_id: str,
    date: str,
    serial: str,
    animals: List[AnimalRow],
    out_path: str,
):
    try:
        from docx import Document
        from docx.shared import Pt, Mm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        messagebox.showerror("Missing Library", "python-docx not installed.\nRun: pip install python-docx")
        return False

    clinic = CLINICS[clinic_key]
    doc = Document()

    section = doc.sections[0]
    section.page_width  = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = section.right_margin = Mm(18)
    section.top_margin = section.bottom_margin = Mm(15)

    def add_ar_para(text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT,
                    color: tuple = None) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.font.name = "Amiri"  # or "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    # Title
    add_ar_para(clinic["name_ar"], bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_ar_para(f'سجل تجاري : {clinic["trade_reg"]}  |  الضريبي : {clinic["tax_reg"]}',
                size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("─" * 70)

    add_ar_para(f'اسم المالك : {owner}   |   رقم الهوية : {owner_id}', bold=True)
    add_ar_para(f'تاريخ الإصدار : {date}', bold=True, color=(192, 57, 43))
    doc.add_paragraph()
    paragraph_text = PARAGRAPH_TEMPLATE.format(clinic["phone"])
    add_ar_para(paragraph_text, size=11)
    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate([
        "العدد", "الجنس", "رقم الشريحة / الحلقة", "العمر", "الاسم العلمي"
    ]):
        hdr[i].text = h
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for a in animals:
        row = table.add_row().cells
        age_txt = a.age  # age already includes unit
        for i, val in enumerate([str(a.count), a.gender, a.id_number, age_txt, a.species]):
            row[i].text = val
            row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_ar_para(f'رقم الإصدار : {serial}', bold=True, color=(192, 57, 43))
    add_ar_para(clinic["address_ar"])
    add_ar_para(f'ترخيص العيادة : {clinic["license"]}')
    doc.add_paragraph()

    # Add stamp and signature images (with larger signature)
    def add_image(path, width_inches=1.5):
        if path and os.path.exists(path):
            try:
                doc.add_picture(path, width=Inches(width_inches))
            except Exception as e:
                print(f"[DOCX] Could not add image {path}: {e}")
                doc.add_paragraph(f"[صورة غير متاحة: {os.path.basename(path)}]")
        else:
            doc.add_paragraph(f"[{os.path.basename(path) if path else 'صورة'} غير موجودة]")

    # Two-column table for stamp and signature
    img_table = doc.add_table(rows=1, cols=2)
    img_table.autofit = False
    img_table.columns[0].width = Inches(2.5)
    img_table.columns[1].width = Inches(2.5)
    left_cell = img_table.rows[0].cells[0]
    right_cell = img_table.rows[0].cells[1]
    left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add stamp and signature to cells only (no duplicate)
    def add_image_to_cell(cell, path, width_inches):
        if path and os.path.exists(path):
            try:
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(path, width=Inches(width_inches))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if cell == left_cell else WD_ALIGN_PARAGRAPH.RIGHT
            except Exception as e:
                cell.text = f"[خطأ في الصورة]"
        else:
            cell.text = f"[لا توجد صورة]"

    add_image_to_cell(left_cell, clinic["stamp_path"], 1.3)
    # Signature larger (2 inches)
    add_image_to_cell(right_cell, SIG_1, 2.0)

    doc.save(out_path)
    return True

# ─────────────────────────────────────────────
#  SPECIES MANAGER WINDOW (unchanged)
# ─────────────────────────────────────────────
class SpeciesManagerWindow(tb.Toplevel):
    def __init__(self, parent, species_list: List[str], on_save):
        super().__init__(parent)
        self.title("إدارة الأسماء العلمية")
        self.geometry("480x540")
        self.resizable(False, False)
        self.species  = list(species_list)
        self.on_save  = on_save
        self._build()

    def _build(self):
        ttk.Label(self, text="الأسماء العلمية المتاحة",
                  font=("Helvetica", 13, "bold")).pack(pady=(12,4))

        frame = ttk.Frame(self)
        frame.pack(fill=BOTH, expand=True, padx=12, pady=4)

        sb = ttk.Scrollbar(frame)
        sb.pack(side=RIGHT, fill=Y)

        self.lb = tk.Listbox(frame, yscrollcommand=sb.set, font=("Helvetica", 10),
                             selectmode=tk.SINGLE, activestyle="dotbox")
        self.lb.pack(fill=BOTH, expand=True)
        sb.config(command=self.lb.yview)
        self._refresh_list()

        add_frame = ttk.Frame(self)
        add_frame.pack(fill=X, padx=12, pady=4)
        self.new_var = tk.StringVar()
        ttk.Entry(add_frame, textvariable=self.new_var, font=("Helvetica", 10)).pack(side=LEFT, fill=X, expand=True, padx=(0,6))
        ttk.Button(add_frame, text="إضافة", bootstyle=SUCCESS,
                   command=self._add).pack(side=LEFT)

        btn_frame = ttk.Frame(self)
        btn_frame.pack(fill=X, padx=12, pady=(0,8))
        ttk.Button(btn_frame, text="حذف المحدد", bootstyle=DANGER,
                   command=self._delete).pack(side=LEFT, padx=(0,8))
        ttk.Button(btn_frame, text="حفظ وإغلاق", bootstyle=PRIMARY,
                   command=self._save).pack(side=RIGHT)

    def _refresh_list(self):
        self.lb.delete(0, tk.END)
        for sp in sorted(self.species):
            self.lb.insert(tk.END, sp)

    def _add(self):
        name = self.new_var.get().strip()
        if not name:
            return
        if name not in self.species:
            self.species.append(name)
            self._refresh_list()
        self.new_var.set("")

    def _delete(self):
        sel = self.lb.curselection()
        if not sel:
            return
        name = self.lb.get(sel[0])
        if messagebox.askyesno("حذف", f"حذف '{name}'?"):
            self.species = [s for s in self.species if s != name]
            self._refresh_list()

    def _save(self):
        save_species(self.species)
        self.on_save(self.species)
        self.destroy()

# ─────────────────────────────────────────────
#  MAIN APPLICATION (unchanged except minor fixes)
# ─────────────────────────────────────────────
class App(tb.Window):
    def __init__(self):
        super().__init__(themename="cosmo")
        self.title("مولّد شهادات الصحة — الأنعام الأليفة")
        self.geometry("1100x720")
        self.minsize(900, 620)
        self.configure(padx=14, pady=10)

        self.species_list: List[str] = load_species()
        self.animals:      List[AnimalRow] = []

        self._build_ui()

    @staticmethod
    def _bind_clip(w: tk.Widget):
        for seq, ev in [("<Control-a>", "<<SelectAll>>"),
                         ("<Control-c>", "<<Copy>>"),
                         ("<Control-v>", "<<Paste>>"),
                         ("<Control-x>", "<<Cut>>")]:
            w.bind(seq, lambda e, ev=ev: w.event_generate(ev))

    def _build_ui(self):
        # Clinic bar
        clinic_bar = ttk.Frame(self)
        clinic_bar.pack(fill=X, pady=(0, 6))
        ttk.Label(clinic_bar, text="العيادة :", font=("Helvetica", 11, "bold")).pack(side=RIGHT, padx=(0,6))
        self.var_clinic = tk.StringVar(value="خميس مشيط")
        for name in CLINICS:
            ttk.Radiobutton(clinic_bar, text=name, variable=self.var_clinic,
                            value=name, bootstyle="primary-toolbutton").pack(side=RIGHT, padx=4)

        # Owner card
        info_card = ttk.LabelFrame(self, text=" بيانات المالك ", padding=10)
        info_card.pack(fill=X, pady=(0, 6))

        self.var_owner  = tk.StringVar()
        self.var_id     = tk.StringVar()
        self.var_date   = tk.StringVar(value=datetime.now().strftime('%Y/%m/%d'))
        self.var_serial = tk.StringVar(value=self._new_serial())

        fields = [
            ("اسم المالك", self.var_owner, 0, 0, 30),
            ("رقم الهوية",  self.var_id,   0, 2, 20),
            ("التاريخ",     self.var_date,  0, 4, 16),
            ("رقم الإصدار", self.var_serial,0, 6, 18),
        ]
        for lbl, var, r, c, w in fields:
            ttk.Label(info_card, text=lbl).grid(row=r, column=c, sticky="e", padx=(8,2), pady=4)
            e = ttk.Entry(info_card, textvariable=var, width=w)
            e.grid(row=r, column=c+1, sticky="ew", padx=(0,10), pady=4)
            self._bind_clip(e)
        info_card.columnconfigure(1, weight=3)
        info_card.columnconfigure(3, weight=2)

        # Animal input card
        mid = ttk.LabelFrame(self, text=" إضافة حيوان ", padding=10)
        mid.pack(fill=X, pady=(0, 6))

        ttk.Label(mid, text="الاسم العلمي").grid(row=0, column=0, sticky="e", padx=(4,2))
        self.var_species = tk.StringVar(value=self.species_list[0] if self.species_list else "")
        self.sci_cb = ttk.Combobox(mid, textvariable=self.var_species,
                                   values=self.species_list, width=26)
        self.sci_cb.grid(row=1, column=0, padx=4, pady=2, sticky="ew")
        self._bind_clip(self.sci_cb)
        ttk.Button(mid, text="⚙ إدارة الأسماء", bootstyle=SECONDARY,
                   command=self._open_species_mgr).grid(row=1, column=1, padx=4)

        ttk.Label(mid, text="العمر").grid(row=0, column=2, sticky="e", padx=(8,2))
        self.var_age = tk.StringVar()
        self.var_age_unit = tk.StringVar(value="شهور")
        age_frame = ttk.Frame(mid)
        age_frame.grid(row=1, column=2, padx=4, pady=2)
        e_age = ttk.Entry(age_frame, textvariable=self.var_age, width=6)
        e_age.pack(side=LEFT, padx=(0,3))
        self._bind_clip(e_age)
        ttk.Radiobutton(age_frame, text="شهور", variable=self.var_age_unit,
                        value="شهور", bootstyle="secondary-toolbutton").pack(side=LEFT, padx=1)
        ttk.Radiobutton(age_frame, text="سنة", variable=self.var_age_unit,
                        value="سنة",  bootstyle="secondary-toolbutton").pack(side=LEFT, padx=1)

        ttk.Label(mid, text="نوع الرقم").grid(row=0, column=3, sticky="e", padx=(8,2))
        self.var_id_type = tk.StringVar(value="شريحة")
        self.var_id_type_custom = tk.StringVar()
        id_type_frame = ttk.Frame(mid)
        id_type_frame.grid(row=1, column=3, padx=4, pady=2)
        ttk.Radiobutton(id_type_frame, text="شريحة", variable=self.var_id_type,
                        value="شريحة", bootstyle="info-toolbutton").pack(side=LEFT, padx=2)
        ttk.Radiobutton(id_type_frame, text="حلقة", variable=self.var_id_type,
                        value="حلقة", bootstyle="info-toolbutton").pack(side=LEFT, padx=2)
        ttk.Radiobutton(id_type_frame, text="أخرى:", variable=self.var_id_type,
                        value="__custom__", bootstyle="info-toolbutton").pack(side=LEFT, padx=2)
        e_custom_type = ttk.Entry(id_type_frame, textvariable=self.var_id_type_custom, width=8)
        e_custom_type.pack(side=LEFT, padx=2)

        ttk.Label(mid, text="الرقم").grid(row=0, column=4, sticky="e", padx=(8,2))
        self.var_chip = tk.StringVar()
        e_chip = ttk.Entry(mid, textvariable=self.var_chip, width=22)
        e_chip.grid(row=1, column=4, padx=4, pady=2, sticky="ew")
        self._bind_clip(e_chip)

        ttk.Label(mid, text="الجنس").grid(row=0, column=5, sticky="e", padx=(8,2))
        self.var_gender = tk.StringVar(value="أنثى")
        ttk.Combobox(mid, textvariable=self.var_gender,
                     values=["ذكر", "أنثى"], width=8).grid(row=1, column=5, padx=4, pady=2)

        btn_frame = ttk.Frame(mid)
        btn_frame.grid(row=1, column=6, padx=(10,0))
        ttk.Button(btn_frame, text="+ إضافة", bootstyle=SUCCESS,
                   command=self._add_animal).pack(side=LEFT, padx=2)
        ttk.Button(btn_frame, text="لصق شرائح", bootstyle=INFO,
                   command=self._paste_dialog).pack(side=LEFT, padx=2)

        mid.columnconfigure(0, weight=2)
        mid.columnconfigure(4, weight=2)

        # Table
        tbl_card = ttk.LabelFrame(self, text=" قائمة الحيوانات ", padding=6)
        tbl_card.pack(fill=BOTH, expand=True, pady=(0, 6))

        cols = ("species", "age", "id_type", "id_number", "gender", "count")
        col_hdrs = ("الاسم العلمي", "العمر", "النوع", "الرقم", "الجنس", "العدد")
        col_widths = (200, 80, 70, 160, 70, 60)

        self.tree = ttk.Treeview(tbl_card, columns=cols, show="headings", height=10)
        for c, h, w in zip(cols, col_hdrs, col_widths):
            self.tree.heading(c, text=h)
            self.tree.column(c, width=w, anchor="center")
        self.tree.pack(side=LEFT, fill=BOTH, expand=True)
        self.tree.bind("<Double-1>", self._edit_row)

        vsb = ttk.Scrollbar(tbl_card, orient=VERTICAL, command=self.tree.yview)
        vsb.pack(side=RIGHT, fill=Y)
        self.tree.configure(yscrollcommand=vsb.set)

        del_btn = ttk.Button(self, text="🗑 حذف المحدد", bootstyle=DANGER,
                             command=self._delete_selected)
        del_btn.pack(anchor="w", pady=(0,4))

        # Bottom
        bot = ttk.Frame(self)
        bot.pack(fill=X)

        self.var_out = tk.StringVar(value=OUT_DIR)
        ttk.Label(bot, text="مجلد الحفظ :").pack(side=LEFT, padx=(0,4))
        eout = ttk.Entry(bot, textvariable=self.var_out, width=40)
        eout.pack(side=LEFT, padx=(0,4), fill=X, expand=True)
        self._bind_clip(eout)
        ttk.Button(bot, text="📁 تغيير", bootstyle=SECONDARY,
                   command=self._choose_dir).pack(side=LEFT, padx=(0,12))

        self.var_word = tk.BooleanVar(value=False)
        ttk.Checkbutton(bot, text="تصدير Word أيضاً", variable=self.var_word,
                        bootstyle="round-toggle").pack(side=LEFT, padx=(0,12))

        ttk.Button(bot, text="🔄 شهادة جديدة", bootstyle=INFO,
                   command=self._new_cert).pack(side=RIGHT, padx=(6,0))
        ttk.Button(bot, text="✅ توليد الشهادة", bootstyle=PRIMARY,
                   command=self._generate).pack(side=RIGHT)

    def _open_species_mgr(self):
        def on_save(new_list):
            self.species_list = new_list
            self.sci_cb["values"] = new_list
        SpeciesManagerWindow(self, self.species_list, on_save)

    def _insert_row(self, a: AnimalRow):
        age_txt = a.age  # age already includes unit (e.g. "6 شهور" / "2 سنة")
        self.tree.insert("", "end",
            values=(a.species, age_txt, a.id_type, a.id_number, a.gender, a.count))

    def _add_animal(self):
        sp  = self.var_species.get().strip()
        _age_raw  = self.var_age.get().strip()
        _age_unit = self.var_age_unit.get()
        ag = f"{_age_raw} {_age_unit}" if _age_raw else ""
        ch  = self.var_chip.get().strip()
        idt = self.var_id_type.get()
        if idt == "__custom__":
            idt = self.var_id_type_custom.get().strip() or "أخرى"
        ge  = self.var_gender.get()

        if not ch:
            messagebox.showerror("خطأ", "يجب إدخال رقم التعريف")
            return

        a = AnimalRow(sp or "Unknown", ag, ch, idt, ge, 1)
        self.animals.append(a)
        self._insert_row(a)
        self.var_chip.set("")

    def _delete_selected(self):
        for sel in self.tree.selection():
            idx = self.tree.index(sel)
            self.tree.delete(sel)
            if 0 <= idx < len(self.animals):
                self.animals.pop(idx)

    def _edit_row(self, event):
        sel = self.tree.selection()
        if not sel:
            return
        item = sel[0]
        idx  = self.tree.index(item)
        vals = self.tree.item(item, "values")

        win = tb.Toplevel(self)
        win.title("تعديل حيوان")
        win.geometry("420x340")
        win.grab_set()

        fields = [
            ("الاسم العلمي", tk.StringVar(value=vals[0]), "combo"),
            ("العمر",        tk.StringVar(value=vals[1]), "entry"),
            ("نوع الرقم",    tk.StringVar(value=vals[2]), "radio"),
            ("الرقم",        tk.StringVar(value=vals[3]), "entry"),
            ("الجنس",        tk.StringVar(value=vals[4]), "gender"),
        ]
        widgets = {}
        for i, (lbl, var, kind) in enumerate(fields):
            ttk.Label(win, text=lbl).grid(row=i, column=0, sticky="e", padx=10, pady=5)
            if kind == "combo":
                w = ttk.Combobox(win, textvariable=var, values=self.species_list, width=28)
            elif kind == "radio":
                f = ttk.Frame(win)
                ttk.Radiobutton(f, text="شريحة", variable=var, value="شريحة").pack(side=LEFT)
                ttk.Radiobutton(f, text="حلقة",  variable=var, value="حلقة").pack(side=LEFT)
                w = f
            elif kind == "gender":
                w = ttk.Combobox(win, textvariable=var, values=["ذكر","أنثى"], width=12)
            else:
                w = ttk.Entry(win, textvariable=var, width=28)
            if isinstance(w, ttk.Frame):
                w.grid(row=i, column=1, sticky="w", padx=10, pady=5)
            else:
                w.grid(row=i, column=1, sticky="ew", padx=10, pady=5)
            widgets[lbl] = var

        def save():
            sp   = widgets["الاسم العلمي"].get()
            age  = widgets["العمر"].get().strip()
            idt  = widgets["نوع الرقم"].get()
            chip = widgets["الرقم"].get().strip()
            ge   = widgets["الجنس"].get()
            self.animals[idx] = AnimalRow(sp, age, chip, idt, ge, self.animals[idx].count)
            age_txt = age  # age already includes unit
            self.tree.item(item, values=(sp, age_txt, idt, chip, ge, self.animals[idx].count))
            win.destroy()

        ttk.Button(win, text="حفظ", bootstyle=SUCCESS, command=save
                   ).grid(row=len(fields), column=0, columnspan=2, pady=12)

    def _paste_dialog(self):
        win = tb.Toplevel(self)
        win.title("لصق أرقام الشرائح / الحلقات")
        win.geometry("500x430")
        win.grab_set()

        ttk.Label(win, text="الصق الأرقام (كل رقم في سطر)").pack(pady=(10,2))
        txt = tk.Text(win, height=8, font=("Courier", 10))
        txt.pack(fill=BOTH, expand=True, padx=10)

        ttk.Label(win, text="الاسم العلمي").pack(pady=(8,0))
        var_sci = tk.StringVar(value=self.var_species.get())
        ttk.Combobox(win, textvariable=var_sci,
                     values=self.species_list, width=34).pack()

        ttk.Label(win, text="العمر").pack(pady=(6,0))
        paste_age_frame = ttk.Frame(win)
        paste_age_frame.pack()
        var_age = tk.StringVar(value=self.var_age.get())
        ttk.Entry(paste_age_frame, textvariable=var_age, width=8).pack(side=LEFT, padx=(0,4))
        var_age_unit_p = tk.StringVar(value=self.var_age_unit.get())
        ttk.Radiobutton(paste_age_frame, text="شهور", variable=var_age_unit_p,
                        value="شهور", bootstyle="secondary-toolbutton").pack(side=LEFT, padx=2)
        ttk.Radiobutton(paste_age_frame, text="سنة", variable=var_age_unit_p,
                        value="سنة",  bootstyle="secondary-toolbutton").pack(side=LEFT, padx=2)

        ttk.Label(win, text="نوع الرقم").pack(pady=(6,0))
        var_idt = tk.StringVar(value=self.var_id_type.get())
        rf = ttk.Frame(win)
        rf.pack()
        ttk.Radiobutton(rf, text="شريحة", variable=var_idt, value="شريحة").pack(side=LEFT, padx=8)
        ttk.Radiobutton(rf, text="حلقة",  variable=var_idt, value="حلقة").pack(side=LEFT)

        gf = ttk.Frame(win)
        gf.pack(fill=X, padx=10, pady=6)
        ttk.Label(gf, text="عدد الذكور:").pack(side=LEFT)
        var_m = tk.IntVar(value=0)
        ttk.Entry(gf, textvariable=var_m, width=5).pack(side=LEFT, padx=4)
        ttk.Label(gf, text="عدد الإناث:").pack(side=LEFT, padx=(12,0))
        var_f = tk.IntVar(value=0)
        ttk.Entry(gf, textvariable=var_f, width=5).pack(side=LEFT, padx=4)

        def apply():
            chips = [ln.strip() for ln in txt.get("1.0","end").splitlines() if ln.strip()]
            if not chips:
                messagebox.showerror("خطأ", "لا توجد أرقام"); return
            m, f = var_m.get(), var_f.get()
            if m + f != len(chips):
                messagebox.showerror("خطأ",
                    f"مجموع الجنسين ({m}+{f}={m+f}) ≠ عدد الأرقام ({len(chips)})"); return
            genders = ["ذكر"]*m + ["أنثى"]*f
            random.shuffle(genders)
            for chip, gender in zip(chips, genders):
                _combined_age = f"{var_age.get().strip()} {var_age_unit_p.get()}" if var_age.get().strip() else ""
                a = AnimalRow(var_sci.get() or "Unknown", _combined_age,
                              chip, var_idt.get(), gender, 1)
                self.animals.append(a)
                self._insert_row(a)
            messagebox.showinfo("تم", f"تمت إضافة {len(chips)} حيوانات")
            win.destroy()

        ttk.Button(win, text="إضافة", bootstyle=SUCCESS, command=apply).pack(pady=8)

    @staticmethod
    def _new_serial() -> str:
        return str(random.randint(10000, 99999))

    def _choose_dir(self):
        d = filedialog.askdirectory(initialdir=self.var_out.get())
        if d:
            self.var_out.set(d)

    def _new_cert(self):
        if not messagebox.askyesno("تأكيد", "مسح كل البيانات وبدء شهادة جديدة؟"):
            return
        self.var_owner.set("")
        self.var_id.set("")
        self.var_date.set(datetime.now().strftime('%Y/%m/%d'))
        self.var_serial.set(self._new_serial())
        self.var_chip.set("")
        self.var_age.set("")
        self.var_age_unit.set("شهور")
        self.animals.clear()
        for item in self.tree.get_children():
            self.tree.delete(item)

    def _generate(self):
        owner  = self.var_owner.get().strip()
        oid    = self.var_id.get().strip()
        date   = self.var_date.get().strip()
        serial = self.var_serial.get().strip()
        clinic = self.var_clinic.get()
        out_d  = self.var_out.get().strip() or OUT_DIR

        if not (owner and oid and serial):
            messagebox.showerror("خطأ", "الرجاء إدخال اسم المالك ورقم الهوية ورقم الإصدار")
            return
        if not self.animals:
            messagebox.showerror("خطأ", "أضف حيواناً واحداً على الأقل")
            return

        os.makedirs(out_d, exist_ok=True)

        first = self.animals[0]
        safe = lambda s: "".join(c for c in s if c.isalnum() or c in " _-").strip()
        base_name = f"{safe(owner)}_{safe(first.species)}_{first.id_type}"
        pdf_path  = os.path.join(out_d, base_name + ".pdf")
        docx_path = os.path.join(out_d, base_name + ".docx")

        try:
            build_pdf(clinic, owner, oid, date, serial, self.animals, pdf_path)
            msg = f"✅ PDF محفوظ:\n{pdf_path}"

            # Save archival copy in output_isdar named by serial number
            isdar_dir = os.path.join(out_d, "output_isdar")
            os.makedirs(isdar_dir, exist_ok=True)
            isdar_path = os.path.join(isdar_dir, f"{serial}.pdf")
            build_pdf(clinic, owner, oid, date, serial, self.animals, isdar_path)
            msg += f"\n✅ نسخة الأرشيف:\n{isdar_path}"

            if self.var_word.get():
                ok = build_docx(clinic, owner, oid, date, serial, self.animals, docx_path)
                if ok:
                    msg += f"\n\n✅ Word محفوظ:\n{docx_path}"

            messagebox.showinfo("تم الحفظ", msg)

        except Exception as e:
            messagebox.showerror("خطأ", f"فشل توليد الشهادة:\n{e}")
            import traceback; traceback.print_exc()

if __name__ == "__main__":
    App().mainloop()